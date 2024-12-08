from fastapi.middleware.cors import CORSMiddleware
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import FileResponse
from typing import Optional, List
from pathlib import Path
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
import os
import json
import logging
import shutil
import stat
from uuid import uuid4
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings
from langchain_core.prompts import PromptTemplate
from utils import *


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR: Optional[Path] = "default"
DB_DIRECTORY:Optional[Path] = os.path.dirname(UPLOAD_DIR)

load_dotenv()
API_KEY = os.getenv('OPENAI_API_KEY')

client = OpenAI(api_key=API_KEY)

#initialize vector database 
embedings = OpenAIEmbeddings(api_key = API_KEY
                             ,model="text-embedding-3-small")
vectore_store = Chroma(embedding_function=embedings
                       ,persist_directory=DB_DIRECTORY)

retriever = vectore_store.as_retriever(search_kwargs = {"k":5})

def add_document_to_vector_db(document_name):
    content = extract_text(document_name)
    new_document = Document(
        page_content= content 
        ,metadata={"name":document_name}
    )

    vectore_store.add_documents([new_document],ids = [str(uuid4())])


def readDocx(file_path):
    doc = Document(file_path)
    res  = ""
    for paragraph in doc.paragraphs:
        res += paragraph.text
    return res

@app.post("/api/search")
async def search_files(request: dict):
    query = request.get("query")
    
    logger.info(f"Starting search with query: {query}")
    logger.info(f"Using API key: {client.api_key[:13]}...")

    try:
        retrieve_answer = retriever.invoke(query)
        
        # Process the response to extract relevant files
        relevant_files = [f.metadata['name']for f in retrieve_answer]
        return {"results": relevant_files}
    except Exception as e:
        print(f"Search error details: {type(e).__name__}: {str(e)}")  # Detailed error logging
        import traceback
        print(traceback.format_exc())  # Print full stack trace
        return {"error": f"OpenAI API error: {str(e)}"}


@app.post("/api/search_assistant")
async def get_assistant_responce(request):
    user_query = request.get("query")
    logger.info(f"Starting search with query: {user_query}")
    logger.info(f"Using API key: {client.api_key[:13]}...")

    try:
        retrieve_data = retriever.invoke(user_query)
        context = retrieve_data[0].page_content +'\n'+retrieve_data[1].page_content
    
        PROMPT_TEMPLATE = PromptTemplate.from_template('''You are a highly skilled Due Diligence Assistant designed to support users in conducting and precise research.
        Use this context {context} to answer {query}.
        Make your answers short.''')
    
        answer = client.chat.completions.create(
        model="gpt-4o"
        ,messages=[
            {
                "role":"user"
                ,"content":PROMPT_TEMPLATE.format(context = context,query = user_query)
            }
        ]   
        )
        return {"results":answer}
    except Exception as e:
        print(f"Search error details: {type(e).__name__}: {str(e)}")  # Detailed error logging
        import traceback
        print(traceback.format_exc())  # Print full stack trace
        return {"error": f"OpenAI API error: {str(e)}"}


@app.post("/api/summarize")
async def summarize_file(request: dict):
    filename = request.get("filename")
    current_file_path = request.get("currentPath")

    if not filename:
        return {"error": "No filename provided"}
        
    try:
        #logger.info(f"Filename to Read content: {filename}")

        if(os.path.splitext(filename)[1] != ".docx"):
            logger.warning("Only DOCX files supported for summarization now")
            return

        # Read file content
        file_path = UPLOAD_DIR / current_file_path / filename

        logger.info(f"Filepath to Read: {file_path}")
        
        content = readDocx(file_path)

        # with open(file_path, 'r') as f:
        #     content = f.readlines()
            
        logger.info(f"content read: {content[:100]}")

        # Get summary from GPT
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a file summarization assistant. Provide a concise summary of the file content."},
                {"role": "user", "content": f"Please summarize this file:\n{content[:5000]}"}
            ]
        )
        
        return {"summary": response.choices[0].message.content}
    except Exception as e:
        return {"error": str(e)}

@app.post("/set-config")
async def set_config(config: dict):
    global UPLOAD_DIR
    try:
        logger.info(f"Received config: {config}")
        
        if "config_path" not in config:
            raise ValueError("config_path not provided")
            
        config_path = config["config_path"]
        logger.info(f"Config path: {config_path}")
        
        if not os.path.exists(config_path):
            raise FileNotFoundError(f"Config file not found at {config_path}")
            
        with open(config_path) as f:
            config_data = json.load(f)
            logger.info(f"Loaded config data: {config_data}")
            
            if "upload_dir" not in config_data:
                raise ValueError("upload_dir not found in config file")
                
            UPLOAD_DIR = Path(config_data["upload_dir"])
            UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
            
            # Set directory permissions to 755 (rwxr-xr-x)
            os.chmod(UPLOAD_DIR, stat.S_IRWXU | stat.S_IRGRP | stat.S_IXGRP | stat.S_IROTH | stat.S_IXOTH)
            
            logger.info(f"Set UPLOAD_DIR to {UPLOAD_DIR} with full permissions")
            
        return {"status": "success", "upload_dir": str(UPLOAD_DIR)}
    except Exception as e:
        logger.error(f"Error setting config: {str(e)}", exc_info=True)
        return {"error": str(e)}

@app.get("/config")
async def get_config():
    """Endpoint to check current configuration"""
    return {
        "upload_dir": str(UPLOAD_DIR) if UPLOAD_DIR else None,
        "upload_dir_exists": UPLOAD_DIR.exists() if UPLOAD_DIR else False if UPLOAD_DIR else None
    }

# Add validation to other endpoints
def validate_upload_dir():
    if UPLOAD_DIR is None:
        return {"error": "Upload directory not configured"}
    if not UPLOAD_DIR.exists():
        return {"error": "Upload directory does not exist"}

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    try:
        # Use the filename from the uploaded file
        filename = file.filename.replace('\\', '/')
        
        # Create full path in upload directory
        full_path = UPLOAD_DIR / filename
        
        # Ensure parent directory exists
        os.makedirs(str(full_path.parent), exist_ok=True)
        
        # Save the file
        with open(full_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        try:
            add_document_to_vector_db(file)
        except Exception as e:
            print("{e}")
            
        return {
            "filename": filename,
            "status": "success"
        }
    except Exception as e:
        return {"error": str(e)}

@app.post("/upload-folder")
async def upload_folder(files: List[UploadFile] = File(...), path: str = Form("")):
    validate_upload_dir()
    try:
        results = []
        for file in files:
            # Create full path including any subdirectories
            full_path = UPLOAD_DIR / path / file.filename
            
            # Ensure the directory structure exists
            full_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Save the file
            with open(full_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            try:
                add_document_to_vector_db(file)
            except Exception as e:
                print("{e}")
            results.append({
                "filename": file.filename,
                "path": str(full_path.relative_to(UPLOAD_DIR)),
                "status": "success"
            })
        
        return {"files": results}
    except Exception as e:
        return {"error": str(e)}

@app.get("/download/{file_path:path}")
async def download_file(file_path: str):
    full_path = UPLOAD_DIR / file_path
    if full_path.exists():
        return FileResponse(full_path)
    return {"error": "File not found"}

@app.delete("/delete/{file_path:path}")
async def delete_file(file_path: str):
    full_path = UPLOAD_DIR / file_path
    try:
        if full_path.is_file():
            os.remove(full_path)
            # Remove empty parent directories
            current_dir = full_path.parent
            while current_dir != UPLOAD_DIR:
                if not any(current_dir.iterdir()):
                    current_dir.rmdir()
                    current_dir = current_dir.parent
                else:
                    break
            return {"status": "success", "message": f"{file_path} deleted"}
        return {"error": "File not found"}
    except Exception as e:
        return {"error": str(e)}

@app.get("/files")
async def list_files(path: str = ""):
    try:
        # Combine UPLOAD_DIR with the requested path
        current_dir = UPLOAD_DIR / path if path else UPLOAD_DIR
        logger.info(f"Listing files in directory: {current_dir}")
        
        entries = []
        if current_dir.exists():
            for entry in current_dir.iterdir():
                # Get path relative to the current directory
                relative_entry = entry.relative_to(current_dir)
                if entry.is_dir():
                    entries.append(f"{str(relative_entry)}/")
                else:
                    entries.append(str(relative_entry))
                    
        logger.info(f"Found entries: {entries}")
        # Sort folders first, then files, both alphabetically
        return {"entries": sorted(entries, key=lambda x: (not x.endswith('/'), x.lower()))}
    except Exception as e:
        logger.error(f"Error listing files: {str(e)}", exc_info=True)
        return {"error": str(e)}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)