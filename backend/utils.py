from docx import Document
from PyPDF2 import PdfReader
from langchain_openai import OpenAIEmbeddings
from langchain_chroma import Chroma
import logging
import os

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def init_retriever(api_key, db_path):
    #initialize vector database 
    embedings = OpenAIEmbeddings(api_key = api_key
                                ,model="text-embedding-3-small")
    vectore_store = Chroma(embedding_function=embedings
                        ,persist_directory=db_path)
    retriever = vectore_store.as_retriever(search_kwargs = {"k":5})
    return (vectore_store, retriever)

def extract_text_from_docx(file_path):
    """Extract text from a .docx file."""
    try:
        logger.info(f"extract_text_from_docx ENTER FilePath: {file_path}")
        # doc = Document(file_path)
        # res  = ""
        # for paragraph in doc.paragraphs:
        #     res += paragraph.text

        #logger.info(f"extract_text_from_docx FINISH Result: {res}")
        #return res
        doc = Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        return f"Error reading .docx file: {e}"

def extract_text_from_pdf(file_path):
    """Extract text from a .pdf file."""
    try:
        reader = PdfReader(file_path)
        return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    except Exception as e:
        return f"Error reading .pdf file: {e}"

def extract_text_from_txt(file_path):
    """Extract text from a .txt file."""
    try:
        with open(file_path, 'r') as file:
            return file.read()
    except Exception as e:
        return f"Error reading .txt file: {e}"

def extract_text(file_path):
    """Determine the file type and extract text accordingly."""
    logger.info(f"extract_text ENTER FilePath: {file_path}")
    if os.path.splitext(file_path)[1] == '.docx':
        return extract_text_from_docx(file_path)
    elif os.path.splitext(file_path)[1] == '.pdf':
        return extract_text_from_pdf(file_path)
    return extract_text_from_txt(file_path)
    

